import sys, os, json, re, base64, zipfile
import xml.etree.ElementTree as ET

WORKDIR = os.getcwd()


def load_input_json():
    p = os.path.join(WORKDIR, "input.json")
    if os.path.exists(p):
        with open(p, "r", encoding="utf-8") as f:
            return json.load(f)
    return {}


def find_docx_in_workdir():
    for fn in os.listdir(WORKDIR):
        if fn.lower().endswith(".docx"):
            return os.path.join(WORKDIR, fn)
    return None


def resolve_docx_path(args, input_data):
    if args:
        candidate = args[0]
        if os.path.exists(candidate):
            return candidate
        candidate2 = os.path.join(WORKDIR, candidate)
        if os.path.exists(candidate2):
            return candidate2
    found = find_docx_in_workdir()
    if found:
        return found
    b64 = input_data.get("docx_base64")
    if b64:
        out_path = os.path.join(WORKDIR, "input_transcript.docx")
        with open(out_path, "wb") as f:
            f.write(base64.b64decode(b64))
        return out_path
    return None


def extract_with_python_docx(path):
    import docx  # python-docx
    d = docx.Document(path)
    paragraphs = [p.text for p in d.paragraphs if p.text and p.text.strip()]
    table_rows = []
    for t in d.tables:
        for row in t.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                table_rows.append(" | ".join(cells))
    return paragraphs, table_rows


W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def extract_with_stdlib(path):
    paragraphs = []
    table_rows = []
    with zipfile.ZipFile(path) as z:
        with z.open("word/document.xml") as f:
            tree = ET.parse(f)
    root = tree.getroot()
    body = root.find(f"{W_NS}body")
    if body is None:
        return paragraphs, table_rows
    for elem in body:
        tag = elem.tag
        if tag == f"{W_NS}p":
            texts = [n.text for n in elem.iter(f"{W_NS}t") if n.text]
            line = "".join(texts).strip()
            if line:
                paragraphs.append(line)
        elif tag == f"{W_NS}tbl":
            for tr in elem.iter(f"{W_NS}tr"):
                cells = []
                for tc in tr.iter(f"{W_NS}tc"):
                    texts = [n.text for n in tc.iter(f"{W_NS}t") if n.text]
                    cells.append("".join(texts).strip())
                if any(cells):
                    table_rows.append(" | ".join(cells))
    return paragraphs, table_rows


SPEAKER_TS_PATTERNS = [
    re.compile(r"^\[?(?P<ts>\d{1,2}:\d{2}(:\d{2})?)\]?\s*[-\u2013]?\s*(?P<speaker>[A-Za-z][\w .'-]{0,60}):\s*(?P<text>.*)$"),
    re.compile(r"^(?P<speaker>[A-Za-z][\w .'-]{0,60}):\s*(?P<text>.*)$"),
    re.compile(r"^\[?(?P<ts>\d{1,2}:\d{2}(:\d{2})?)\]?\s*(?P<text>.*)$"),
]


def segment(paragraphs):
    segments = []
    speakers = set()
    for i, line in enumerate(paragraphs):
        matched = False
        for pat in SPEAKER_TS_PATTERNS:
            m = pat.match(line)
            if m:
                gd = m.groupdict()
                speaker = gd.get("speaker")
                ts = gd.get("ts")
                text = gd.get("text", line).strip()
                segments.append({"index": i, "speaker": speaker, "timestamp": ts, "text": text})
                if speaker:
                    speakers.add(speaker)
                matched = True
                break
        if not matched:
            segments.append({"index": i, "speaker": None, "timestamp": None, "text": line})
    return segments, speakers


def main():
    args = sys.argv[1:]
    input_data = load_input_json()
    output_name = input_data.get("output_name", "transcript")

    docx_path = resolve_docx_path(args, input_data)
    if not docx_path:
        print(json.dumps({"error": "No .docx file found. Provide inputFiles, inputData.docx_base64, or an args path."}))
        sys.exit(1)

    warnings = []
    try:
        paragraphs, table_rows = extract_with_python_docx(docx_path)
        engine = "python-docx"
    except Exception as e:
        warnings.append(f"python-docx unavailable or failed ({e}); used stdlib XML fallback")
        try:
            paragraphs, table_rows = extract_with_stdlib(docx_path)
            engine = "stdlib-xml"
        except Exception as e2:
            print(json.dumps({"error": f"Failed to parse docx with both engines: {e2}"}))
            sys.exit(1)

    all_lines = paragraphs + ([f"[TABLE] {r}" for r in table_rows] if table_rows else [])
    segments, speakers = segment(all_lines)

    txt_path = os.path.join(WORKDIR, f"{output_name}.txt")
    with open(txt_path, "w", encoding="utf-8") as f:
        f.write("\n".join(all_lines))

    json_path = os.path.join(WORKDIR, f"{output_name}_segments.json")
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(segments, f, indent=2, ensure_ascii=False)

    summary = {
        "engine_used": engine,
        "paragraph_count": len(paragraphs),
        "table_row_count": len(table_rows),
        "segment_count": len(segments),
        "speakers_detected": sorted(speakers),
        "warnings": warnings,
        "output_files": [os.path.basename(txt_path), os.path.basename(json_path)],
    }

    print(json.dumps(summary, indent=2))
    print("----- TRANSCRIPT TEXT START -----")
    print("\n".join(all_lines))
    print("----- TRANSCRIPT TEXT END -----")


if __name__ == "__main__":
    main()
